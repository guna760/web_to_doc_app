import io
import requests
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def create_docx(data):
    doc = Document()
    doc.add_heading(data['title'], level=0)
    
    for item in data['content']:
        if item['type'] == 'text':
            doc.add_paragraph(item['value'])
        elif item['type'] == 'image':
            try:
                img_response = requests.get(item['value'], timeout=5)
                if img_response.status_code == 200:
                    image_stream = io.BytesIO(img_response.content)
                    doc.add_paragraph().add_run().add_picture(image_stream, width=Inches(5.5))
            except Exception:
                continue
                
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def create_pdf(data):
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    story = []
    
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=18, spaceAfter=20)
    story.append(Paragraph(data['title'], title_style))
    story.append(Spacer(1, 15))
    
    for item in data['content']:
        if item['type'] == 'text':
            clean_text = item['value'].encode('utf-8', 'ignore').decode('utf-8')
            story.append(Paragraph(clean_text, normal_style))
            story.append(Spacer(1, 10))
        elif item['type'] == 'image':
            try:
                img_response = requests.get(item['value'], timeout=5)
                if img_response.status_code == 200:
                    image_stream = io.BytesIO(img_response.content)
                    rl_img = RLImage(image_stream, width=400, height=250)
                    story.append(rl_img)
                    story.append(Spacer(1, 10))
            except Exception:
                continue
        
    doc.build(story)
    bio.seek(0)
    return bio